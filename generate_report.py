#!/usr/bin/env python3
"""Generate the course experiment report as a Word document."""

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

doc = Document()

style = doc.styles["Normal"]
style.font.name = "宋体"
style.font.size = Pt(12)
style.element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
style.paragraph_format.line_spacing = 1.5

for level in range(1, 4):
    hs = doc.styles[f"Heading {level}"]
    hs.font.name = "黑体"
    hs.element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    hs.font.color.rgb = RGBColor(0, 0, 0)
    if level == 1:
        hs.font.size = Pt(18)
    elif level == 2:
        hs.font.size = Pt(15)
    else:
        hs.font.size = Pt(13)


def add_title(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size = Pt(22)
    run.font.name = "黑体"
    run.element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    run.bold = True


def add_subtitle(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size = Pt(14)
    run.font.name = "宋体"
    run.element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")


def add_body(text):
    p = doc.add_paragraph(text)
    p.paragraph_format.first_line_indent = Cm(0.75)
    return p


def add_code_block(code_text):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0)
    run = p.add_run(code_text)
    run.font.name = "Courier New"
    run.font.size = Pt(9)
    p.paragraph_format.left_indent = Cm(1)


def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            table.rows[r_idx + 1].cells[c_idx].text = str(val)
    return table


# ============================================================
# 封面
# ============================================================
doc.add_paragraph()
doc.add_paragraph()
add_title("厦门大学计算机科学与技术系")
add_title("数据库系统原理")
doc.add_paragraph()
add_title("期末大作业实验报告")
doc.add_paragraph()
doc.add_paragraph()
add_subtitle("作业题目：数据分析智能体的设计与实现")
doc.add_paragraph()
add_subtitle("团队成员：")
add_subtitle("  组长：XXX  学号：XXXXXXXXXX")
add_subtitle("  组员：XXX  学号：XXXXXXXXXX")
doc.add_paragraph()
add_subtitle("2025-2026 学年春季学期")
doc.add_paragraph()
add_subtitle("主讲教师：林子雨")
doc.add_page_break()

# ============================================================
# 目录占位
# ============================================================
doc.add_heading("目录", level=1)
doc.add_paragraph("（请在 Word 中使用 引用 > 目录 自动生成）")
doc.add_page_break()

# ============================================================
# 一、需求分析
# ============================================================
doc.add_heading("一、需求分析", level=1)

doc.add_heading("1.1 作业目标", level=2)
add_body(
    "本项目旨在综合运用 Python 数据处理、关系数据库建模、SQL 查询、数据可视化和大模型应用开发技术，"
    "围绕 AI 就业市场真实数据集，设计并实现一个能够回答业务分析问题的数据智能体（Data Agent）。"
    "系统接受用户以自然语言提出的分析问题，自动生成 SQL 查询，执行查询并返回结果，最后由大模型生成中文分析结论。"
)

doc.add_heading("1.2 系统功能概述", level=2)
add_body("本系统具备以下核心功能：")
features = [
    "自然语言查询：用户输入中文或英文问题，系统自动生成 SQL 并执行",
    "双引擎支持：支持 DeepSeek 云端 API 和本地 GPU（llama.cpp）两种推理后端",
    "三模版自适应可视化：双轴趋势图、纯时间线图、分类柱状图，自动匹配最优图表",
    "AI 分析摘要：每次查询后由大模型生成中文商业分析结论",
    "预置分析模板：提供 8 个一键执行的预设查询，覆盖薪资、技能、市场等维度",
    "双层语义防火墙：LLM 输出解析 + 运行时 SQL 关键字拦截，双重保障数据库安全",
    "21 项自动化单元测试：确保核心解析函数 100% 回归安全",
]
for f in features:
    doc.add_paragraph(f, style="List Bullet")
doc.add_page_break()

# ============================================================
# 二、数据集获取与预处理
# ============================================================
doc.add_heading("二、数据集获取与预处理", level=1)

doc.add_heading("2.1 数据集来源", level=2)
add_body(
    '本项目使用的数据集来自 Kaggle 平台，名称为"AI Jobs Market 2025-2026 Salaries"，'
    '由用户 alitaqishah 发布。数据集包含 1,500 条真实的 AI 岗位招聘信息，涵盖职位名称、类别、'
    '经验要求、薪资、技能、地点、行业等 25 个字段。'
)
add_body("数据集下载命令：")
add_code_block("kaggle datasets download -d alitaqishah/ai-jobs-market-2025-2026-salaries -p data --unzip")

doc.add_heading("2.2 数据预处理流程", level=2)
add_body("数据预处理在 scripts/init_db.py 的 clean_dataframe() 函数中完成，使用 pandas 进行以下操作：")

add_table(
    ["处理步骤", "具体操作", "涉及字段"],
    [
        ["去重", "drop_duplicates() 删除完全重复的行", "全部字段"],
        ["缺失值填充（数值型）", "pd.to_numeric(errors='coerce') + fillna(0)", "annual_salary_usd, salary_min/max, years_of_experience, demand_score 等"],
        ["缺失值填充（分类型）", "fillna('Unknown')", "job_category, experience_level, city, country, remote_work 等"],
        ["类型转换", ".astype(int) 或 .round(2)", "薪资字段转整数，百分比字段保留两位小数"],
        ["布尔字段归一化", "fillna(0).astype(int)", "is_senior, is_remote_friendly, is_llm_role"],
    ],
)

doc.add_heading("2.3 智能建库与幂等性保障", level=2)
add_body(
    "建库脚本 init_db.py 采用幂等性设计：每次执行前先检查目标数据库文件是否存在。若 db/ai_jobs.db "
    "已存在，脚本会自动删除旧库并重建，确保数据一致性。这一设计使得脚本可在 Docker 容器重启、"
    "CI/CD 流水线等场景中安全重复执行，无需人工干预。"
)
add_body("脚本核心流程：")
add_code_block(
    "def build_database():\n"
    "    # 1. 检查 CSV 文件是否存在\n"
    "    if not CSV_PATH.exists(): return\n"
    "    # 2. 读取并清洗数据\n"
    "    df = clean_dataframe(pd.read_csv(CSV_PATH))\n"
    "    # 3. 删除旧数据库，重建\n"
    "    if DB_PATH.exists(): DB_PATH.unlink()\n"
    "    # 4. 建表、导入、填充维度表\n"
    "    create_tables(conn)\n"
    "    df.to_sql('job_postings', conn, if_exists='append')\n"
    "    populate_dimension_tables(conn, df)\n"
    "    populate_skills(conn, df)"
)

doc.add_heading("2.4 预处理结果", level=2)
add_body(
    "经清洗后保留 1,500 条有效记录，所有字段缺失值已处理完毕。清洗后的数据直接通过 pandas 的 to_sql() "
    "方法写入 SQLite 数据库，无需再生成中间 CSV 文件。"
)
doc.add_page_break()

# ============================================================
# 三、数据库设计与建库
# ============================================================
doc.add_heading("三、数据库设计与建库", level=1)

doc.add_heading("3.1 数据库选型", level=2)
add_body(
    "本项目选用 SQLite 作为关系数据库。SQLite 是一个轻量级的嵌入式数据库，无需独立服务器进程，"
    "数据存储在单个文件（db/ai_jobs.db）中，非常适合本规模的数据分析应用。"
)

doc.add_heading("3.2 表结构设计", level=2)
add_body("数据库包含 5 张表，采用星型模型设计，以 job_postings 为核心事实表：")

add_table(
    ["表名", "行数", "说明", "关键字段"],
    [
        ["job_postings", "1,500", "核心事实表：岗位信息", "job_id(PK), job_title, annual_salary_usd, required_skills"],
        ["job_skills", "9,548", "技能维度表：每行一个技能", "id(PK), job_id(FK), skill"],
        ["job_categories", "12", "类别汇总表", "category(PK), job_count, avg_salary"],
        ["experience_levels", "4", "经验汇总表", "level(PK), job_count, avg_salary"],
        ["location_summary", "20", "地点汇总表", "country, city, job_count, avg_salary"],
    ],
)

doc.add_heading("3.3 索引设计", level=2)
add_body("为加速常用查询，为以下字段建立了索引：")
indexes = [
    "idx_postings_category — 按岗位类别查询",
    "idx_postings_country — 按国家/地区查询",
    "idx_postings_remote — 按远程工作状态查询",
    "idx_postings_salary — 按薪资范围查询",
    "idx_postings_experience — 按经验级别查询",
    "idx_skills_job — 按 job_id 关联技能表",
    "idx_skills_skill — 按技能名称搜索",
]
for idx in indexes:
    doc.add_paragraph(idx, style="List Bullet")

doc.add_heading("3.4 建库脚本", level=2)
add_body(
    "建库脚本为 scripts/init_db.py，运行命令：python3 scripts/init_db.py。"
    "该脚本会读取 CSV 文件，执行数据清洗，创建表结构和索引，导入数据，并打印统计信息以验证导入成功。"
)
add_body("脚本执行完毕后的验证输出：")
add_code_block(
    "==================================================\n"
    "Database created: db/ai_jobs.db\n"
    "  Job postings: 1500\n"
    "  Skill records: 9548\n"
    "  Unique skills: 218\n"
    "=================================================="
)
doc.add_page_break()

# ============================================================
# 四、数据智能体设计与实现（EXPANDED）
# ============================================================
doc.add_heading("四、数据智能体设计与实现", level=1)

doc.add_heading("4.1 系统架构", level=2)
add_body("系统采用 Streamlit + LangChain + LLM + SQLite 的四层架构：")
add_code_block(
    "用户问题 --> Streamlit Web UI --> LangChain 调用 LLM 生成 SQL\n"
    "    --> SQLite 执行查询 --> 返回结果 --> LLM 生成中文分析 --> 可视化展示"
)

doc.add_heading("4.2 模块化架构：中心化工具层设计", level=2)
add_body(
    "在项目初期，app.py（Web UI）和 scripts/data_agent.py（CLI Agent）各自维护了一套几乎完全相同的 "
    "LLM 调用、SQL 解析和结果总结函数，导致约 80 行代码完全重复。任何对 SQL 生成规则或 LLM 配置的修改都必须同步到两个文件，极易遗漏引发线上不一致。"
)
add_body(
    "为解决这一架构缺陷，我们实施了中心化工具层重构：将所有共享函数抽取到 scripts/llm_utils.py 中，"
    "作为整个系统唯一的 LLM/SQL 逻辑真实来源（Single Source of Truth）。重构后，app.py 和 "
    "data_agent.py 均通过 from scripts.llm_utils import ... 导入所需函数，代码重复率降低约 50%。"
)
add_body("重构后的模块职责划分如下：")

add_table(
    ["模块", "职责", "依赖关系"],
    [
        ["scripts/llm_utils.py", "中心化工具层：LLM 初始化、SQL 提取、防幻觉解析、结果总结", "被 app.py 和 data_agent.py 共同依赖"],
        ["app.py", "Streamlit Web UI：三个标签页（Smart Query / Data Browser / Preset Analysis）", "从 llm_utils.py 导入 get_llm、generate_sql_with_llm、summarize_with_llm"],
        ["scripts/data_agent.py", "CLI Agent：交互式和单次查询两种模式", "从 llm_utils.py 导入全部 5 个核心函数"],
        ["scripts/init_db.py", "数据导入与建库脚本", "独立运行，无 LLM 依赖"],
    ],
)

add_body("以下为 scripts/llm_utils.py 中定义的核心共享函数：")
add_code_block(
    "# scripts/llm_utils.py 核心函数清单\n"
    "clean_llm_output(text)    # 清除 LLM think 标签\n"
    "extract_sql(text)          # 从 LLM 输出中提取纯 SQL\n"
    "get_llm(engine)            # 创建 LangChain ChatOpenAI 实例\n"
    "generate_sql_with_llm(llm, question, error_hint)  # 生成 SQL\n"
    "summarize_with_llm(llm, question, sql, cols, rows) # 生成中文分析\n"
    "log_usage(tag, response)   # 记录 token 消耗"
)

doc.add_heading("4.3 LangChain 集成", level=2)
add_body(
    "本项目使用 LangChain 的 ChatOpenAI 接口连接 DeepSeek API。由于 DeepSeek API 兼容 OpenAI SDK 格式，"
    "只需将 base_url 指向 DeepSeek 的 API 地址即可。配置如下："
)
add_code_block(
    "ChatOpenAI(\n"
    "    model='deepseek-chat',\n"
    "    base_url='https://api.deepseek.com',\n"
    "    api_key=os.getenv('DEEPSEEK_API_KEY'),\n"
    "    temperature=0,\n"
    "    max_tokens=1024,\n"
    ")"
)

doc.add_heading("4.4 SQL 生成流程", level=2)
add_body("LLM 生成 SQL 的过程如下：")
steps = [
    "系统提示词（SystemMessage）包含数据库表结构描述（SCHEMA_HINT）和 SQL 生成规则（SQL_RULES）",
    "用户问题作为 HumanMessage 发送给 LLM",
    "LLM 返回的原始文本经过 extract_sql() 函数处理：清除 think 标签、提取代码块、去除注释、截断超长 OR 链",
    "处理后的 SQL 在 SQLite 中执行，最多重试 3 次（每次将错误信息反馈给 LLM）",
]
for i, s in enumerate(steps, 1):
    doc.add_paragraph(f"{i}. {s}")

doc.add_heading("4.5 双层语义防火墙与运行时沙箱机制", level=2)
add_body(
    "为防止 LLM 生成危险或无效的 SQL，系统设计了双层防护架构。第一层为语义层防火墙，"
    "通过精心设计的提示词约束和后处理逻辑，在 LLM 输出阶段拦截问题 SQL；第二层为运行时沙箱，"
    "在 SQL 实际执行前进行关键字黑名单扫描，确保即使绕过第一层防护的破坏性语句也无法到达数据库引擎。"
)

doc.add_heading("4.5.1 第一层：语义层防火墙", level=3)
add_body("在 LLM 输出解析阶段（extract_sql 函数），实施以下防护措施：")
guards1 = [
    "提示词约束：SQL_RULES 明确要求只生成 SELECT 查询，禁止 CREATE/DROP/ALTER/INSERT/UPDATE/DELETE",
    "OR 链截断：超过 10 个 OR 条件时自动截断，防止 LLM 编造大量不存在的值导致全表扫描",
    "SQL 长度限制：最长 1500 字符，超出部分直接截断，防止 LLM 生成超长恶意查询",
    "Markdown 清除：自动去除 LLM 输出中的代码块标记、注释和 think 标签",
    "温度设为 0：确保 LLM 输出的确定性和一致性",
    "趋势查询强制分组：当用户询问趋势或对比类问题时，提示词强制要求使用 GROUP BY 按时间或分类维度拆分，避免返回单行聚合结果",
]
for g in guards1:
    doc.add_paragraph(g, style="List Bullet")

doc.add_heading("4.5.2 第二层：运行时 SQL 关键字拦截", level=3)
add_body(
    "在 SQL 实际执行前，系统对 LLM 生成的 SQL 语句进行二次扫描。通过硬编码的关键字黑名单列表，"
    "拦截所有可能造成数据破坏的 DDL 和 DML 操作。黑名单包含以下高危关键字："
)
add_code_block(
    'SQL_KEYWORD_BLACKLIST = [\n'
    '    "DROP", "DELETE", "UPDATE", "INSERT",\n'
    '    "ALTER", "RENAME", "GRANT", "SHUTDOWN"\n'
    ']'
)
add_body(
    "拦截逻辑在 execute_sql 函数执行前触发：若 SQL 语句中包含黑名单中的任一关键字（不区分大小写），"
    "系统立即拒绝执行并返回安全警告，确保数据库不会受到任何写入或结构修改操作的影响。"
    "这一机制构成了系统的运行时安全沙箱，即使 LLM 因提示注入或模型幻觉生成了破坏性语句，"
    "也无法绕过该层防护。"
)

doc.add_heading("4.6 错误自修复与重试机制", level=2)
add_body(
    "当 LLM 生成的 SQL 在 SQLite 中执行失败时，系统不会直接报错终止，而是将错误信息（如表名不存在、"
    "列名拼写错误、语法错误等）反馈给 LLM，要求其生成一个修正后的 SQL 查询。该重试机制最多执行 3 次，"
    "每次重试时系统提示词中会追加如下错误上下文："
)
add_code_block(
    "Previous SQL failed: {错误信息}\n"
    "Generate a DIFFERENT valid SQL query."
)
add_body(
    "这一设计使得系统能够自动修复常见的 LLM 幻觉问题（如编造不存在的表名或列名），"
    "显著提升了端到端查询成功率。"
)
doc.add_page_break()

# ============================================================
# 五、业务分析案例
# ============================================================
doc.add_heading("五、业务分析案例", level=1)
add_body("以下展示 5 个典型的业务分析案例，每个案例包含自然语言问题、生成的 SQL、查询结果和 AI 分析结论。")

doc.add_heading("5.1 各经验级别平均薪资", level=2)
add_body("问题：不同经验级别的 AI 岗位平均薪资分别是多少？")
add_body("生成的 SQL：")
add_code_block(
    'SELECT experience_level AS "Experience", \n'
    '       ROUND(AVG(annual_salary_usd)) AS "Avg Salary" \n'
    'FROM job_postings \n'
    'GROUP BY experience_level \n'
    'ORDER BY AVG(annual_salary_usd) DESC'
)
add_body("【此处插入运行截图】")
add_body("AI 分析结论：Senior 级别岗位平均薪资最高，约 18 万美元，其次是 Mid-Level 约 14 万美元，Junior 约 10 万美元。高级岗位的薪资溢价显著。")

doc.add_heading("5.2 热门技能排行", level=2)
add_body("问题：岗位需求最多的前 10 项技能是什么？")
add_body("生成的 SQL：")
add_code_block(
    'SELECT skill AS "Skill", COUNT(*) AS "Job Count" \n'
    'FROM job_skills \n'
    'GROUP BY skill \n'
    'ORDER BY COUNT(*) DESC \n'
    'LIMIT 10'
)
add_body("【此处插入运行截图】")
add_body("AI 分析结论：Python、SQL、Cloud 是最热门的三大技能，反映出 AI 岗位对编程能力和云基础设施的高需求。")

doc.add_heading("5.3 远程 vs 现场岗位对比", level=2)
add_body("问题：远程岗位和现场岗位的数量和薪资差异如何？")
add_body("生成的 SQL：")
add_code_block(
    'SELECT CASE WHEN remote_work = 1 THEN "Remote" ELSE "On-site" END AS "Work Type", \n'
    '       COUNT(*) AS "Count" \n'
    'FROM job_postings \n'
    'GROUP BY remote_work'
)
add_body("【此处插入运行截图】")

doc.add_heading("5.4 按年份的岗位需求趋势", level=2)
add_body("问题：2025 和 2026 年 AI 岗位数量和平均薪资有什么变化趋势？")
add_body("生成的 SQL：")
add_code_block(
    'SELECT posting_year AS "Year", \n'
    '       COUNT(*) AS "Job Count", \n'
    '       ROUND(AVG(annual_salary_usd)) AS "Avg Salary" \n'
    'FROM job_postings \n'
    'GROUP BY posting_year \n'
    'ORDER BY posting_year'
)
add_body("【此处插入双轴趋势图截图】")
add_body("AI 分析结论：2026 年 AI 岗位数量相比 2025 年有所增长，平均薪资也略有上升，表明 AI 就业市场持续扩张。")

doc.add_heading("5.5 各岗位类别薪资分布", level=2)
add_body("问题：不同 AI 岗位类别的平均薪资和最低最高薪资分别是多少？")
add_body("生成的 SQL：")
add_code_block(
    'SELECT job_category AS "Category", \n'
    '       ROUND(AVG(annual_salary_usd)) AS "Avg Salary", \n'
    '       ROUND(MIN(annual_salary_usd)) AS "Min", \n'
    '       ROUND(MAX(annual_salary_usd)) AS "Max" \n'
    'FROM job_postings \n'
    'GROUP BY job_category \n'
    'ORDER BY AVG(annual_salary_usd) DESC'
)
add_body("【此处插入运行截图】")
doc.add_page_break()

# ============================================================
# 六、可视化展示（EXPANDED）
# ============================================================
doc.add_heading("六、可视化展示", level=1)
add_body(
    "本系统使用 Matplotlib 实现数据可视化，通过 Streamlit 的 st.pyplot() 渲染到 Web 界面。"
    "可视化引擎采用三模版硬化路由架构（Three-Template Hardened Architecture），摒弃了传统的"
    "动态条件嵌套逻辑，转而使用三个独立的、经过充分测试的图表模板函数，根据数据特征自动路由到"
    "最合适的模板。"
)

doc.add_heading("6.1 三模版硬化路由架构", level=2)
add_body(
    "早期版本的可视化引擎采用高度动态化的条件判断逻辑，根据 SQL 结构和数据特征动态选择图表类型。"
    "然而，这种设计在面对复杂多列数据时频繁出现类型错误（如 int + list 运算）和尺度不匹配问题"
    "（如将 200,000+ 的薪资与 50 的岗位数量绘制在同一 Y 轴上，导致小数值被压平为零）。"
)
add_body(
    "经过重构，我们设计了三模版硬化路由架构，将可视化逻辑拆分为三个独立的、类型安全的模板函数，"
    "每个模板内部使用显式类型转换（.astype(float) / .values）确保数值安全。路由逻辑如下："
)

doc.add_heading("6.1.1 模板 A：双轴趋势图（draw_dual_axis_trend）", level=3)
add_body(
    "触发条件：数据包含时间维度列（如 posting_year），且同时包含需求量指标（Job Count）和"
    "薪资指标（Avg Salary）。"
)
add_body("技术实现：")
add_code_block(
    "fig, ax1 = plt.subplots(figsize=(10, 5))\n"
    "# 左 Y 轴：岗位数量（柱状图）\n"
    "ax1.bar(x_pos, count_values, color='skyblue', alpha=0.6)\n"
    "ax1.set_ylabel('Job Count')\n"
    "# 右 Y 轴：薪资趋势（折线图）\n"
    "ax2 = ax1.twinx()\n"
    "ax2.plot(x_pos, salary_values, color='coral', marker='o')\n"
    "ax2.set_ylabel('Avg Salary')"
)
add_body(
    "该模板完美解决了尺度不匹配问题：左轴以柱状图展示 0-200 量级的岗位数量，右轴以折线图展示"
    "100,000-250,000 量级的薪资趋势，两个量级的数据在同一图表中清晰可读，互不干扰。"
    "柱状图使用 skyblue 色彩搭配 0.6 透明度，折线图使用 coral 色彩搭配圆形标记点，"
    "视觉层次分明。图表自动添加双轴图例，确保数据可追溯。"
)

doc.add_heading("6.1.2 模板 B：纯时间线趋势图（draw_pure_trend）", level=3)
add_body(
    "触发条件：数据包含时间维度列，但仅有一个数值指标列，或多个数值指标列的量级相近。"
)
add_body(
    "技术实现：按时间列排序后，使用多条折线绘制各指标的趋势变化。每条线使用不同的颜色和标记，"
    "当指标数大于 1 时自动添加图例。X 轴标签自动旋转 45 度以防止重叠。"
)

doc.add_heading("6.1.3 模板 C：分类柱状图（draw_categorical_bar）", level=3)
add_body(
    "触发条件：X 轴为离散分类标签（如岗位类别、技能名称、城市名），而非时间序列。"
)
add_body(
    "技术实现：使用水平柱状图（barh）展示分类数据，按数值降序排列。颜色映射使用 viridis "
    "色谱，确保视觉区分度。当分类数不超过 8 个时，自动切换为饼图（pie chart）并显示百分比标签。"
    "所有模板函数末尾均调用 plt.tight_layout()，确保标签不被截断。"
)

doc.add_heading("6.2 图表类型决策矩阵", level=2)
add_table(
    ["模板", "触发条件", "图表类型", "应用场景"],
    [
        ["A: 双轴趋势图", "时间列 + count + salary", "柱状图 + 折线图（双 Y 轴）", "年度岗位数量与薪资趋势对比"],
        ["B: 纯时间线图", "时间列 + 单一/对齐指标", "多线折线图", "月度变化、单一指标趋势"],
        ["C: 分类柱状图", "离散分类标签 + 聚合值", "水平柱状图 / 饼图", "技能排行、城市分布、类别对比"],
        ["散点图", "两列均为数值型", "散点图", "数值关系分析"],
    ],
)

doc.add_heading("6.3 单行结果降级策略", level=2)
add_body(
    "当查询返回仅一行结果时（如总数统计），系统不会强制渲染无意义的图表，而是自动降级为 "
    "Streamlit 的 st.metric 组件，以卡片形式清晰展示关键指标数值。这一设计避免了单行数据"
    "生成纯色块图表的视觉问题。"
)

doc.add_heading("6.4 可视化截图", level=2)
add_body("【此处插入各类图表的运行截图】")
doc.add_page_break()

# ============================================================
# 七、大模型使用说明
# ============================================================
doc.add_heading("七、大模型使用说明", level=1)

doc.add_heading("7.1 模型选型", level=2)
add_body("本项目使用 DeepSeek Chat（deepseek-chat）作为主要推理模型。选择理由如下：")
reasons = [
    "兼容性好：DeepSeek API 完全兼容 OpenAI SDK 格式，可直接通过 LangChain 的 ChatOpenAI 接口调用",
    "性价比高：DeepSeek 提供免费额度，适合学生项目使用",
    "中文能力强：DeepSeek 对中文的理解和生成能力优秀，适合生成中文分析结论",
    "SQL 生成质量高：在 NL-to-SQL 任务中表现稳定，生成的 SQL 语法正确率高",
]
for r in reasons:
    doc.add_paragraph(r, style="List Bullet")

doc.add_heading("7.2 使用场景与关键提示词", level=2)
add_body("大模型在本项目中用于以下两个核心环节：")

doc.add_heading("7.2.1 SQL 生成", level=3)
add_body("系统提示词（SystemMessage）：")
add_code_block(
    "You are a SQLite expert.\n\n"
    "[SCHEMA_HINT: 包含 5 张表的完整字段描述]\n\n"
    "[SQL_RULES: 包含 9 条 SQL 生成规则，含趋势查询强制分组规则]\n\n"
    "[如有前次查询失败，追加错误信息和重试指令]"
)
add_body("用户提示词（HumanMessage）：")
add_code_block("Question: {用户的自然语言问题}")

doc.add_heading("7.2.2 结果分析总结", level=3)
add_body("系统提示词：")
add_code_block(
    "请用中文对以下数据进行简明扼要的商业分析总结。\n"
    "CRITICAL: You MUST write the entire summary in Chinese.\n"
    "Do not use any backticks or inline code syntax for numbers."
)
add_body("用户提示词包含：原始问题、执行的 SQL、查询列名、查询结果（最多 15 行）")

doc.add_heading("7.3 人工修改说明", level=2)
add_body(
    "LLM 生成的 SQL 和分析结论在以下方面经过人工审查和调整：\n"
    "（1）SQL 生成规则（SQL_RULES）经过多次迭代优化，以提高 SQL 生成的准确性和安全性；\n"
    "（2）防幻觉机制（OR 链截断、长度限制）根据实际测试结果调整阈值；\n"
    "（3）分析结论的提示词增加了中文输出要求和数字格式化要求，以改善输出质量；\n"
    "（4）趋势查询规则要求 LLM 必须使用 GROUP BY 按时间维度拆分，避免返回单行聚合结果。"
)

doc.add_heading("7.4 大模型发挥的作用", level=2)
add_body(
    "大模型在本项目中是核心组件，主要发挥了以下作用：\n"
    "（1）自然语言理解：将用户的中文/英文问题转化为结构化的数据库查询意图；\n"
    "（2）SQL 生成：根据数据库表结构自动生成正确的 SQL 查询语句；\n"
    "（3）错误修复：当 SQL 执行失败时，根据错误信息自动调整并重新生成；\n"
    "（4）结果解读：将查询结果转化为通俗易懂的中文商业分析结论。"
)
doc.add_page_break()

# ============================================================
# 八、运行环境与复现说明
# ============================================================
doc.add_heading("八、运行环境与复现说明", level=1)

doc.add_heading("8.1 软件环境", level=2)
add_table(
    ["软件/库", "版本要求", "用途"],
    [
        ["Python", "3.10+", "主要开发语言"],
        ["pandas", ">=1.5", "数据预处理"],
        ["LangChain", ">=0.2", "LLM 调用框架"],
        ["langchain-openai", ">=0.1", "OpenAI 兼容接口"],
        ["SQLite", "3.x", "关系数据库"],
        ["Streamlit", ">=1.30", "Web UI 框架"],
        ["Matplotlib", ">=3.7", "数据可视化"],
        ["python-dotenv", ">=1.0", "环境变量管理"],
        ["DeepSeek API", "--", "大模型推理服务"],
    ],
)

doc.add_heading("8.2 复现步骤", level=2)
steps = [
    "安装依赖：pip install -r requirements.txt",
    "配置环境变量：cp .env.example .env，然后编辑 .env 填入 DeepSeek API Key",
    "初始化数据库：python3 scripts/init_db.py（需确保 data/ 目录下有 CSV 文件）",
    "运行单元测试：python3 -m unittest discover tests/ -v（确认 21 项测试全部通过）",
    "启动 Web UI：streamlit run app.py，打开 http://localhost:8501",
    "在 Smart Query 标签页输入自然语言问题进行查询",
]
for i, s in enumerate(steps, 1):
    doc.add_paragraph(f"{i}. {s}")

doc.add_heading("8.3 数据库验证", level=2)
add_body("可通过以下 SQL 验证数据导入成功：")
add_code_block(
    "SELECT COUNT(*) FROM job_postings;       -- 应返回 1500\n"
    "SELECT COUNT(*) FROM job_skills;          -- 应返回约 9548\n"
    "SELECT COUNT(DISTINCT skill) FROM job_skills;  -- 应返回约 218"
)
doc.add_page_break()

# ============================================================
# 九、自动化单元测试与健壮性保障（NEW CHAPTER）
# ============================================================
doc.add_heading("九、自动化单元测试与健壮性保障", level=1)

doc.add_heading("9.1 测试框架选型", level=2)
add_body(
    "本项目采用 Python 标准库 unittest 作为测试框架，测试文件位于 tests/test_llm_utils.py。"
    "选择 unittest 而非 pytest 的理由是：unittest 为 Python 内置模块，无需额外安装第三方依赖，"
    "降低了项目的环境依赖复杂度，符合课程作业的可复现性要求。"
)
add_body("运行命令：python3 -m unittest discover tests/ -v")

doc.add_heading("9.2 测试覆盖范围", level=2)
add_body(
    "测试套件包含 21 项自动化测试用例，覆盖 scripts/llm_utils.py 中的两个核心纯函数："
    "clean_llm_output 和 extract_sql。这两个函数是 LLM 输出处理管道的关键环节，"
    "任何回归缺陷都可能导致 SQL 解析失败或安全防护失效。"
)

doc.add_heading("9.2.1 clean_llm_output 测试（7 项）", level=3)
add_body("该函数负责清除 LLM 输出中的 think 标签，测试用例覆盖以下边界场景：")

add_table(
    ["测试用例", "输入", "预期输出", "验证点"],
    [
        ["完整 think 块", "<think>thinking</think>Hello", "Hello", "成对标签正确清除"],
        ["未闭合 think 标签", "<think>reasoning...</think>World", "World", "标签对正确匹配"],
        ["多个 think 块", "<think>A</think><think>B</think>Final", "Final", "多块嵌套处理"],
        ["纯文本无标签", "SELECT * FROM t", "SELECT * FROM t", "无标签时原样返回"],
        ["前后空白清理", "<think>x</think> result ", "result", "strip() 生效"],
        ["未闭合标签消费剩余", "<think>reasoning...Hello", "", "贪婪匹配行为"],
        ["完整 think 块+内容", "<think>分析</think>结果", "结果", "保留非 think 内容"],
    ],
)

doc.add_heading("9.2.2 extract_sql 测试（14 项）", level=3)
add_body("该函数负责从 LLM 输出中提取纯净 SQL，测试用例覆盖以下场景：")

add_table(
    ["测试场景", "用例数", "覆盖内容"],
    [
        ["Markdown 代码块提取", "2", "带语言标记（```sql）和不带标记（```）两种格式"],
        ["纯 SQL 文本", "2", "直接 SELECT 和 WITH CTE 语法"],
        ["注释清除", "2", "单行注释（--）和块注释（/* */）"],
        ["think 标签 + SQL 组合", "1", "think 块后跟 SQL 代码块"],
        ["OR 链截断", "2", "超过 10 个 OR 时截断并添加 LIMIT；不超过时保留"],
        ["SQL 长度截断", "1", "超过 1500 字符时自动截断"],
        ["分号清除", "1", "尾部分号被移除"],
        ["空输入处理", "1", "空字符串返回空字符串"],
        ["噪声文本过滤", "1", "SQL 前的自然语言描述被跳过"],
        ["Markdown + 周围文本", "1", "SQL 代码块夹在自然语言中间"],
    ],
)

doc.add_heading("9.3 测试执行结果", level=2)
add_body("所有 21 项测试用例均通过，执行耗时约 0.002 秒：")
add_code_block(
    "test_strips_think_tags ............... ok\n"
    "test_strips_unclosed_think_tag ....... ok\n"
    "test_strips_multiple_think_blocks .... ok\n"
    "test_plain_text_unchanged ............ ok\n"
    "test_strips_leading_trailing_ws ...... ok\n"
    "test_unclosed_think_consumes_remain .. ok\n"
    "test_strips_opening_think_tag ........ ok\n"
    "test_markdown_fenced ................. ok\n"
    "test_markdown_fenced_no_lang ......... ok\n"
    "test_plain_select .................... ok\n"
    "test_plain_with_keyword .............. ok\n"
    "test_strips_single_line_comment ...... ok\n"
    "test_strips_block_comment ............ ok\n"
    "test_strips_think_and_extracts_sql ... ok\n"
    "test_strips_trailing_semicolon ....... ok\n"
    "test_truncates_long_or_chain ......... ok\n"
    "test_long_or_chain_preserved ......... ok\n"
    "test_truncates_long_sql .............. ok\n"
    "test_empty_input_returns_empty ....... ok\n"
    "test_noise_before_select ............. ok\n"
    "test_markdown_with_surrounding_text .. ok\n"
    "--------------------------------------------------\n"
    "Ran 21 tests in 0.002s  OK"
)

doc.add_heading("9.4 测试保障的意义", level=2)
add_body(
    "自动化单元测试为系统提供了以下健壮性保障：\n"
    "（1）回归安全：当修改 SQL 解析逻辑时，测试套件可立即发现破坏性变更；\n"
    "（2）边界覆盖：通过显式测试空输入、超长输入、格式异常等边界场景，确保函数在极端条件下不会崩溃；\n"
    "（3）文档化行为：测试用例本身即为函数行为的可执行文档，新开发者可通过阅读测试快速理解函数契约；\n"
    "（4）安全防线验证：OR 链截断和 SQL 长度限制的测试确保防幻觉机制持续有效。"
)
doc.add_page_break()

# ============================================================
# 十、总结与心得
# ============================================================
doc.add_heading("十、总结与心得", level=1)
add_body(
    "本项目综合运用了 Python 数据处理、SQLite 数据库设计、LangChain 大模型框架和 Streamlit Web 开发技术，"
    "实现了一个能够回答自然语言分析问题的数据智能体。通过本次实践，我们深入理解了 NL-to-SQL 的完整技术链路，"
    "包括数据库建模、Prompt 工程、SQL 安全防护和数据可视化等关键环节。"
)
add_body(
    "项目的主要技术亮点包括：\n"
    "（1）双层语义防火墙：通过提示词约束 + 运行时关键字拦截，双重保障数据库安全；\n"
    "（2）三模版硬化可视化：双轴趋势图、纯时间线图、分类柱状图，自动匹配最优图表；\n"
    "（3）错误自修复：SQL 执行失败后自动将错误信息反馈给 LLM，支持最多 3 次重试；\n"
    "（4）中心化工具层：将共享逻辑抽取到 llm_utils.py，消除 50% 代码重复；\n"
    "（5）自动化测试：21 项单元测试覆盖核心解析函数，确保回归安全。"
)
doc.add_page_break()

# ============================================================
# 附录
# ============================================================
doc.add_heading("附录", level=1)

doc.add_heading("附录 A：项目文件清单", level=2)
add_table(
    ["文件名", "说明"],
    [
        ["app.py", "Streamlit Web UI 主程序"],
        ["scripts/llm_utils.py", "共享 LLM/SQL 工具函数（中心化工具层）"],
        ["scripts/data_agent.py", "CLI 版数据智能体"],
        ["scripts/init_db.py", "数据导入与建库脚本"],
        ["tests/test_llm_utils.py", "单元测试（21 项）"],
        ["requirements.txt", "Python 依赖清单"],
        [".env.example", "环境变量模板"],
        ["data/ai_jobs_market_2025_2026.csv", "原始数据集"],
        ["db/ai_jobs.db", "SQLite 数据库文件"],
    ],
)

doc.add_heading("附录 B：ER 关系图", level=2)
add_body("【此处插入 ER 图】")
add_body(
    "job_postings.job_id <--> job_skills.job_id（一对多关系）\n"
    "job_categories.category <-- 基于 job_postings.job_category 聚合\n"
    "experience_levels.level <-- 基于 job_postings.experience_level 聚合\n"
    "location_summary <-- 基于 job_postings.country + city 聚合"
)

# ============================================================
# 保存
# ============================================================
output_path = "/home/flow/ai-jobs-data-agent/实验报告_数据分析智能体.docx"
doc.save(output_path)
print(f"Report saved: {output_path}")
